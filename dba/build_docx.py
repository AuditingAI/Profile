"""
Build Research_Paper_YMalik_v4.docx from the markdown master.
Handles: front-matter title page, H1-H4, paragraphs with **bold**/*italic*,
bullet & numbered lists, > shaded callout boxes, ![FIGURE: x] images,
and --- page breaks. Body is double-spaced Times New Roman 12 (DBA style).
"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "Research_Paper_YMalik_v4_master.md"
OUT = "Research_Paper_YMalik_v4.docx"

doc = Document()

# ---- base style ----
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
pf.space_after = Pt(0)

for sec in doc.sections:
    sec.top_margin = Inches(1); sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1); sec.right_margin = Inches(1)

def shade(cell, hexfill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hexfill)
    tcPr.append(sh)

def add_inline(p, text):
    # split on **bold** and *italic*
    for tok in re.split(r"(\*\*.+?\*\*|\*[^*]+?\*)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            r = p.add_run(tok[1:-1]); r.italic = True
        else:
            p.add_run(tok)

def single(p):
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

# ---- read & split front matter ----
raw = open(SRC, encoding="utf-8").read().split("\n")
meta = {}
body_start = 0
if raw[0].strip() == "---":
    i = 1
    while i < len(raw) and raw[i].strip() != "---":
        if ":" in raw[i]:
            k, v = raw[i].split(":", 1)
            meta[k.strip()] = v.strip()
        i += 1
    body_start = i + 1

# ---- title page ----
def center(text, size, bold=False, italic=False, after=6):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    single(p); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.italic = italic
    return p

doc.add_paragraph("\n")
center("FLORIDA INTERNATIONAL UNIVERSITY", 14, bold=True)
center("College of Business — Doctor of Business Administration", 12)
doc.add_paragraph("\n")
center(meta.get("TITLE", ""), 18, bold=True, after=10)
doc.add_paragraph("\n")
center(meta.get("AUTHOR", ""), 13, bold=True)
center(meta.get("PROGRAM", ""), 12)
center(meta.get("COURSE", ""), 12)
doc.add_paragraph("\n\n")
center(meta.get("VERSION", ""), 12, italic=True)
center(meta.get("DATE", ""), 12, italic=True)
if meta.get("NOTE"):
    doc.add_paragraph("\n")
    np = doc.add_paragraph(); single(np); np.paragraph_format.space_after = Pt(0)
    r = np.add_run("Drafting note. "); r.italic = True; r.bold = True; r.font.size = Pt(10)
    r2 = np.add_run(meta["NOTE"]); r2.italic = True; r2.font.size = Pt(10)
doc.add_page_break()

# ---- body parsing ----
lines = raw[body_start:]
i = 0
n = len(lines)

def flush_box(buf):
    if not buf:
        return
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    shade(cell, "FFF8E1")
    cell.paragraphs[0].text = ""
    first = True
    for bl in buf:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        single(p); p.paragraph_format.space_after = Pt(2)
        add_inline(p, bl)
        first = False
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

while i < n:
    line = lines[i].rstrip("\n")
    s = line.strip()

    if s.startswith("<!--"):
        i += 1; continue
    if s == "":
        i += 1; continue

    # page break on horizontal rule
    if s == "---":
        doc.add_page_break(); i += 1; continue

    # figure
    m = re.match(r"!\[FIGURE:\s*(.+?)\]", s)
    if m:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            doc.add_picture(m.group(1).strip(), width=Inches(6.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            p.add_run(f"[figure: {m.group(1)} — {e}]")
        i += 1; continue

    # headings
    if s.startswith("#### "):
        doc.add_heading(s[5:], level=4); i += 1; continue
    if s.startswith("### "):
        doc.add_heading(s[4:], level=3); i += 1; continue
    if s.startswith("## "):
        doc.add_heading(s[3:], level=2); i += 1; continue
    if s.startswith("# "):
        doc.add_heading(s[2:], level=1); i += 1; continue

    # blockquote box (collect consecutive)
    if s.startswith(">"):
        buf = []
        while i < n and lines[i].strip().startswith(">"):
            t = lines[i].strip()[1:].strip()
            buf.append(t)
            i += 1
        flush_box(buf)
        continue

    # bullet
    if s.startswith("- "):
        p = doc.add_paragraph(style="List Bullet"); add_inline(p, s[2:]); i += 1; continue

    # numbered
    mnum = re.match(r"^(\d+)\.\s+(.*)", s)
    if mnum:
        p = doc.add_paragraph(style="List Number"); add_inline(p, mnum.group(2)); i += 1; continue

    # reference-pool entries start with "- " already handled; plain paragraph
    p = doc.add_paragraph()
    add_inline(p, s)
    i += 1

doc.save(OUT)
print("wrote", OUT)
print("paragraphs:", len(doc.paragraphs))
