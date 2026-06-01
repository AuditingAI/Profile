"""
Build ALL_IN_ONE_DBA_Package_v1.docx — consolidated master containing every
working file in the system, in one navigable Word document.

Structure:
  Title page
  Table of contents
  PART I  Research Paper (v4.1 complete)
    Ch.1 Introduction (from v3)
    Ch.2 Literature Review (v4.1 expanded)
    Ch.3 Research Model & Hypotheses (v4.1 expanded)
    Ch.4 Methodology (from v3)
    Ch.5 Project Timeline (from v3)
    Ch.6 Recommendations for Future Research (from v3)
    Appendix A Measurement Instrument (draft)
    Verified Reference Pool
  PART II  Supporting Workflow Materials
    Weekly Update W01 v2
    LinkedIn Outreach Pack
    Read Me First
  PART III  Working Toolkit
    NotebookLM Playbook (full)
    Repository & Branch links
"""
import re, json, zipfile, html
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "ALL_IN_ONE_DBA_Package_v3.docx"

master  = open("Research_Paper_YMalik_v4_master.md", encoding="utf-8").read()
playbook = open("NotebookLM_Playbook.md", encoding="utf-8").read()
v3 = json.load(open("/tmp/v3_chapters.json"))

def extract_docx(path):
    z = zipfile.ZipFile(path)
    xml = z.read("word/document.xml").decode("utf-8")
    xml = xml.replace("</w:p>", "\n\n").replace("</w:tr>", "\n")
    text = re.sub(r"<[^>]+>", "", xml)
    text = html.unescape(text)
    return re.sub(r"\n{3,}", "\n\n", text).strip()

weekly   = extract_docx("Weekly_Update_W01_v2.docx")
linkedin = extract_docx("LinkedIn_Outreach_Pack.docx")
readme   = extract_docx("Read_Me_First.docx")
readiness = open("Data_Collection_Readiness.md", encoding="utf-8").read()
rey_2026_06_01 = open("correspondence/2026-06-01_Rey_feedback_and_response.md", encoding="utf-8").read()

# READY packet (filed 2026-06-01 in 03_Recruitment_and_Pilot/)
pilot_readme = open("03_Recruitment_and_Pilot/README.md", encoding="utf-8").read()
pilot_protocol = extract_docx("03_Recruitment_and_Pilot/YMalik_Informed_Pilot_Protocol_READY_2026-06-01.docx")
cloudresearch = extract_docx("03_Recruitment_and_Pilot/YMalik_CloudResearch_Launch_Draft_READY_2026-06-01.docx")
mturk = extract_docx("03_Recruitment_and_Pilot/YMalik_MTurk_Backup_Launch_Draft_READY_2026-06-01.docx")

# Extract xlsx feedback workbook as text tables
def extract_xlsx(path):
    from openpyxl import load_workbook
    wb = load_workbook(path, data_only=True)
    out = []
    for sn in wb.sheetnames:
        ws = wb[sn]
        out.append(f"## Sheet: {sn}\n")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(c.strip() for c in cells):
                out.append(" | ".join(cells))
        out.append("")
    return "\n".join(out)

pilot_xlsx = extract_xlsx("03_Recruitment_and_Pilot/YMalik_Pilot_Feedback_Form_and_Revision_Log_READY_2026-06-01.xlsx")

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(11)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
pf.line_spacing = 1.15
pf.space_after = Pt(4)
for sec in doc.sections:
    sec.top_margin = Inches(0.9); sec.bottom_margin = Inches(0.9)
    sec.left_margin = Inches(1);  sec.right_margin = Inches(1)

def shade(cell, hexfill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hexfill)
    tcPr.append(sh)

def inline(p, text):
    for tok in re.split(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)", text):
        if not tok: continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            r = p.add_run(tok[1:-1]); r.italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(10)
        else:
            p.add_run(tok)

def single(p): p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
def center(text, size, bold=False, italic=False, after=4):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    single(p); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.italic = italic

# ---- TITLE PAGE ----
doc.add_paragraph("\n")
center("FLORIDA INTERNATIONAL UNIVERSITY", 14, bold=True)
center("College of Business  Doctor of Business Administration", 11)
doc.add_paragraph()
center("DBA RESEARCH PROJECT  ALL-IN-ONE PACKAGE", 11, italic=True)
center("(Complete research paper + supporting workflow materials + study toolkit)", 10, italic=True)
doc.add_paragraph()
center("Mitigating Anchoring Bias in", 20, bold=True, after=2)
center("Long-Term Auditor Engagements:", 20, bold=True, after=2)
center("An EFA-Based Validation Study", 18, bold=True, after=8)
doc.add_paragraph()
center("Yasir A. Malik", 13, bold=True)
center("PID 1687105  |  Cohort 7.16  |  Course GEB7913", 11)
center("Instructor / Chair: Professor Juan Rey", 11)
doc.add_paragraph()
center("Compiled: May 30, 2026", 11, italic=True)
center("Package version: ALL-IN-ONE v1 (paper v4.1)", 11, italic=True)
doc.add_paragraph()
np = doc.add_paragraph(); single(np); np.paragraph_format.space_after = Pt(0)
r = np.add_run("Note. "); r.italic = True; r.bold = True; r.font.size = Pt(10)
r2 = np.add_run(
    "This single document consolidates every working file in the DBA project repo as of the compile "
    "date. It is the author's master reference, not the submission to the advisor. The advisor "
    "receives only Part I (the research paper itself). Parts II-III support development."
); r2.italic = True; r2.font.size = Pt(10)
doc.add_page_break()

# ---- TOC ----
doc.add_heading("Contents", level=1)
toc = [
    ("PART I  RESEARCH PAPER (v4.1, complete)", 1),
    ("Chapter 1  Introduction and Statement of the Problem", 2),
    ("Chapter 2  Review of the Literature  (expanded, v4.1)", 2),
    ("Chapter 3  Research Model and Hypotheses  (expanded, v4.1)", 2),
    ("Chapter 4  Methodology", 2),
    ("Chapter 5  Project Timeline (Summary)", 2),
    ("Chapter 6  Recommendations for Future Research", 2),
    ("Appendix A  Measurement Instrument (draft for review)", 2),
    ("Verified Reference Pool", 2),
    ("PART II  ADVISOR CORRESPONDENCE & DATA-COLLECTION READINESS", 1),
    ("2026-06-01  Dr. Rey feedback + response draft", 2),
    ("Data Collection Readiness checklist", 2),
    ("PART III  INFORMED PILOT & RECRUITMENT (READY packet, 2026-06-01)", 1),
    ("03_Recruitment_and_Pilot folder README", 2),
    ("Informed Pilot Protocol", 2),
    ("Pilot Feedback Form + Revision Log (workbook contents)", 2),
    ("CloudResearch Launch Draft (primary route)", 2),
    ("MTurk Backup Launch Draft (secondary route)", 2),
    ("PART IV  SUPPORTING WORKFLOW MATERIALS", 1),
    ("Weekly Update  Week 1 (revised)", 2),
    ("LinkedIn Outreach Pack (IRB-approved recruitment copy)", 2),
    ("Read Me First (folder guide)", 2),
    ("PART V  WORKING TOOLKIT", 1),
    ("NotebookLM Playbook (ADD-friendly extraction workflow)", 2),
    ("Repository and Branch Links", 2),
]
for label, level in toc:
    p = doc.add_paragraph(); single(p); p.paragraph_format.space_after = Pt(2)
    indent = "    " * (level - 1)
    r = p.add_run(f"{indent} {label}")
    if level == 1: r.bold = True; r.font.size = Pt(12)
    else: r.font.size = Pt(11)
doc.add_page_break()

def flush_box(buf):
    if not buf: return
    tbl = doc.add_table(rows=1, cols=1); tbl.style = "Table Grid"
    cell = tbl.cell(0, 0); shade(cell, "FFF8E1")
    cell.paragraphs[0].text = ""
    first = True
    for bl in buf:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        single(p); p.paragraph_format.space_after = Pt(2)
        inline(p, bl); first = False
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

def render_markdown(md, *, drop_front_matter=True):
    lines = md.split("\n")
    i = 0
    if drop_front_matter and lines and lines[0].strip() == "---":
        i = 1
        while i < len(lines) and lines[i].strip() != "---": i += 1
        i += 1
    in_code = False
    code_buf = []
    n = len(lines)
    while i < n:
        line = lines[i].rstrip("\n")
        s = line.strip()
        if s.startswith("```"):
            if not in_code:
                in_code = True; code_buf = []
            else:
                tbl = doc.add_table(rows=1, cols=1); tbl.style = "Table Grid"
                cell = tbl.cell(0, 0); shade(cell, "F4F4F4"); cell.paragraphs[0].text = ""
                first = True
                for cl in code_buf:
                    p = cell.paragraphs[0] if first else cell.add_paragraph()
                    single(p); p.paragraph_format.space_after = Pt(0)
                    r = p.add_run(cl); r.font.name = "Consolas"; r.font.size = Pt(9)
                    first = False
                doc.add_paragraph().paragraph_format.space_after = Pt(0)
                in_code = False
            i += 1; continue
        if in_code:
            code_buf.append(line); i += 1; continue
        if s.startswith("<!--"): i += 1; continue
        if s == "": i += 1; continue
        if s == "---":
            doc.add_page_break(); i += 1; continue
        m = re.match(r"!\[FIGURE:\s*(.+?)\]", s)
        if m:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try: doc.add_picture(m.group(1).strip(), width=Inches(6.4))
            except Exception as e: p.add_run(f"[figure: {m.group(1)}  {e}]")
            i += 1; continue
        if s.startswith("#### "): doc.add_heading(s[5:], level=4); i += 1; continue
        if s.startswith("### "):  doc.add_heading(s[4:], level=3); i += 1; continue
        if s.startswith("## "):   doc.add_heading(s[3:], level=2); i += 1; continue
        if s.startswith("# "):    doc.add_heading(s[2:], level=1); i += 1; continue
        if s.startswith(">"):
            buf = []
            while i < n and lines[i].strip().startswith(">"):
                buf.append(lines[i].strip()[1:].strip()); i += 1
            flush_box(buf); continue
        if s.startswith("- ") or s.startswith("* "):
            p = doc.add_paragraph(style="List Bullet"); inline(p, s[2:]); i += 1; continue
        mnum = re.match(r"^(\d+)\.\s+(.*)", s)
        if mnum:
            p = doc.add_paragraph(style="List Number"); inline(p, mnum.group(2)); i += 1; continue
        if "|" in s and i+1 < n and re.match(r"^\s*\|?\s*[-:|\s]+\|", lines[i+1]):
            rows = []
            while i < n and "|" in lines[i].strip():
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells); i += 1
                if i < n and re.match(r"^\s*\|?\s*[-:|\s]+\|", lines[i]):
                    i += 1
            if rows:
                cols = max(len(r) for r in rows)
                tbl = doc.add_table(rows=len(rows), cols=cols); tbl.style = "Light Grid Accent 1"
                for ri, row in enumerate(rows):
                    for ci in range(cols):
                        c = tbl.cell(ri, ci)
                        txt = row[ci] if ci < len(row) else ""
                        c.text = ""
                        p = c.paragraphs[0]; single(p); inline(p, txt)
                doc.add_paragraph().paragraph_format.space_after = Pt(0)
            continue
        p = doc.add_paragraph(); inline(p, s); i += 1

def emit_plain_text(blob, smart_headings=False):
    for para in blob.split("\n\n"):
        s = para.strip()
        if not s: continue
        if smart_headings:
            if re.match(r"^\d+(\.\d+)*\.?\s+[A-Z][^\n]{0,80}$", s) and len(s) < 90:
                doc.add_heading(s, level=3); continue
            if s.isupper() and len(s) < 70:
                doc.add_heading(s.title(), level=3); continue
        p = doc.add_paragraph(); inline(p, s)

# PART I
doc.add_heading("PART I  RESEARCH PAPER (v4.1, complete)", level=1)
p = doc.add_paragraph()
inline(p, "This part is the complete research paper, combining the v3 baseline (Chapters 1, 4, 5, 6) "
          "with the v4.1 expanded Chapters 2 and 3, Appendix A (instrument), and the Verified Reference "
          "Pool. This is the version to send to the advisor.")
doc.add_page_break()

doc.add_heading("Chapter 1  Introduction and Statement of the Problem", level=2)
emit_plain_text(re.sub(r"^1\.\s+Introduction and Statement of the Problem\s*", "", v3["ch1"]),
                smart_headings=True)
doc.add_page_break()

# v4.1 master already provides Ch.2, Ch.3, Appendix A, References — render it whole
render_markdown(master)
doc.add_page_break()

doc.add_heading("Chapter 4  Methodology", level=2)
emit_plain_text(re.sub(r"^4\.\s+Methodology\s*", "", v3["ch4"]), smart_headings=True)
doc.add_page_break()

doc.add_heading("Chapter 5  Project Timeline (Summary)", level=2)
emit_plain_text(re.sub(r"^5\.\s+Project Timeline.*?\n", "", v3["ch5"]), smart_headings=True)
doc.add_page_break()

doc.add_heading("Chapter 6  Recommendations for Future Research", level=2)
emit_plain_text(re.sub(r"^6\.\s+Recommendations.*?\n", "", v3["ch6"]), smart_headings=True)
doc.add_page_break()

# PART II
doc.add_heading("PART II  ADVISOR CORRESPONDENCE & DATA-COLLECTION READINESS", level=1)
p = doc.add_paragraph()
inline(p, "Verbatim advisor feedback, the author's response draft, and the readiness checklist that "
          "tracks every blocker between current state and live pilot launch. This part is the working "
          "command-and-control surface for the next two weeks.")
doc.add_page_break()

doc.add_heading("2026-06-01  Dr. Rey feedback + response draft", level=2)
render_markdown(rey_2026_06_01, drop_front_matter=False)
doc.add_page_break()

doc.add_heading("Data Collection Readiness checklist", level=2)
render_markdown(readiness, drop_front_matter=False)
doc.add_page_break()

# PART III  Informed Pilot & Recruitment (READY packet)
doc.add_heading("PART III  INFORMED PILOT & RECRUITMENT (READY packet, 2026-06-01)", level=1)
p = doc.add_paragraph()
inline(p, "Filed at dba/03_Recruitment_and_Pilot/ in response to Dr. Rey's 2026-06-01 request to "
          "(a) load the instrument into Qualtrics for face validity, (b) run an informed pilot, and "
          "(c) position the survey on a recruiting platform. The protocol, feedback form, revision log, "
          "and both recruiting-platform launch drafts are all READY.")
doc.add_page_break()

doc.add_heading("03_Recruitment_and_Pilot folder README", level=2)
render_markdown(pilot_readme, drop_front_matter=False)
doc.add_page_break()

doc.add_heading("Informed Pilot Protocol", level=2)
emit_plain_text(pilot_protocol, smart_headings=True)
doc.add_page_break()

doc.add_heading("Pilot Feedback Form + Revision Log (workbook contents)", level=2)
p = doc.add_paragraph()
inline(p, "Workbook with four tabs: Pilot Participant Log, Pilot Feedback Form (PF1-PF10), "
          "Revision Log, and Pilot Decision Summary. The structures are reproduced below as text "
          "tables; use the live .xlsx for actual logging.")
emit_plain_text(pilot_xlsx, smart_headings=True)
doc.add_page_break()

doc.add_heading("CloudResearch Launch Draft (primary route)", level=2)
emit_plain_text(cloudresearch, smart_headings=True)
doc.add_page_break()

doc.add_heading("MTurk Backup Launch Draft (secondary route)", level=2)
emit_plain_text(mturk, smart_headings=True)
doc.add_page_break()

# PART IV  Supporting Workflow Materials
doc.add_heading("PART IV  SUPPORTING WORKFLOW MATERIALS", level=1)
p = doc.add_paragraph()
inline(p, "Course-administrative artifacts and recruitment copy that the paper depends on but does not "
          "include. Do not submit these to the advisor as part of the research paper; they live here "
          "for your reference and as a single record of the project state.")
doc.add_page_break()

doc.add_heading("Weekly Update  Week 1 (revised May 27, 2026)", level=2)
emit_plain_text(weekly, smart_headings=True)
doc.add_page_break()

doc.add_heading("LinkedIn Outreach Pack (IRB-approved recruitment copy)", level=2)
emit_plain_text(linkedin, smart_headings=True)
doc.add_page_break()

doc.add_heading("Read Me First (folder guide)", level=2)
emit_plain_text(readme, smart_headings=True)
doc.add_page_break()

# PART V
doc.add_heading("PART V  WORKING TOOLKIT", level=1)
p = doc.add_paragraph()
inline(p, "Tools the author uses to extend the paper. The NotebookLM Playbook is the primary workflow "
          "for adding more verified citations per construct without losing focus.")
doc.add_page_break()

doc.add_heading("NotebookLM Playbook  DBA Paper Material Extraction", level=2)
playbook_body = re.sub(r"^# .+?\n", "", playbook, count=1)
render_markdown(playbook_body, drop_front_matter=False)
doc.add_page_break()

doc.add_heading("Repository and Branch Links", level=2)
links = [
    ("All-in-One package (this document)",
     "https://github.com/AuditingAI/Profile/blob/claude/scholar-links-review-Plgk6/dba/ALL_IN_ONE_DBA_Package_v3.docx"),
    ("Interactive dashboard (open in browser)",
     "https://github.com/AuditingAI/Profile/blob/claude/scholar-links-review-Plgk6/dba/dashboard.html"),
    ("Research paper v4.1 (send this to Prof. Rey)",
     "https://github.com/AuditingAI/Profile/blob/claude/scholar-links-review-Plgk6/dba/Research_Paper_YMalik_v4.docx"),
    ("Research paper markdown master (editable source)",
     "https://github.com/AuditingAI/Profile/blob/claude/scholar-links-review-Plgk6/dba/Research_Paper_YMalik_v4_master.md"),
    ("Research model figure (PNG)",
     "https://github.com/AuditingAI/Profile/blob/claude/scholar-links-review-Plgk6/dba/Anchoring_Bias_Research_Model.png"),
    ("NotebookLM Playbook (clickable on GitHub)",
     "https://github.com/AuditingAI/Profile/blob/claude/scholar-links-review-Plgk6/dba/NotebookLM_Playbook.md"),
    ("Original v3 proposal (baseline for diff)",
     "https://github.com/AuditingAI/Profile/blob/claude/scholar-links-review-Plgk6/dba/Final_Proposal_YMalik_v3.docx"),
    ("Branch folder (all files)",
     "https://github.com/AuditingAI/Profile/tree/claude/scholar-links-review-Plgk6/dba"),
    ("Commit history",
     "https://github.com/AuditingAI/Profile/commits/claude/scholar-links-review-Plgk6"),
]
for label, url in links:
    p = doc.add_paragraph(); single(p); p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f" {label}: "); r1.bold = True; r1.font.size = Pt(10)
    r2 = p.add_run(url); r2.font.size = Pt(10); r2.font.name = "Consolas"

doc.save(OUT)
print("wrote", OUT)
print("paragraphs:", len(doc.paragraphs))
print("tables:", len(doc.tables))
