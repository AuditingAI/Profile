"""Build the W06 weekly-update docx from its markdown (front-matter title block,
H1-H4, **bold**/*italic*/`code`, bullets, numbered lists, --- page breaks)."""
import re, sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

SRC = sys.argv[1] if len(sys.argv) > 1 else "Weekly_Update_W06_2026-06-21_FINAL.md"
OUT = SRC.rsplit(".", 1)[0] + ".docx"

doc = Document()
n = doc.styles["Normal"]
n.font.name = "Times New Roman"; n.font.size = Pt(11)
n._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
n.paragraph_format.line_spacing = 1.15
n.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
n.paragraph_format.space_after = Pt(6)
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.9)
    s.left_margin = s.right_margin = Inches(1)

def inline(p, text):
    for tok in re.split(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)", text):
        if not tok: continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            r = p.add_run(tok[1:-1]); r.italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1]); r.font.name = "Courier New"
        else:
            p.add_run(tok)

raw = open(SRC, encoding="utf-8").read().split("\n")
meta, body_start = {}, 0
if raw[0].strip() == "---":
    i = 1
    while i < len(raw) and raw[i].strip() != "---":
        if ":" in raw[i]:
            k, v = raw[i].split(":", 1); meta[k.strip()] = v.strip()
        i += 1
    body_start = i + 1

for line in raw[body_start:]:
    s = line.strip()
    if s == "":
        continue
    if s == "---":
        doc.add_paragraph().add_run().add_break(); continue
    if s.startswith("#### "): doc.add_heading(s[5:], level=4); continue
    if s.startswith("### "):  doc.add_heading(s[4:], level=3); continue
    if s.startswith("## "):   doc.add_heading(s[3:], level=2); continue
    if s.startswith("# "):    doc.add_heading(s[2:], level=1); continue
    if s.startswith("> "):
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
        inline(p, s[2:]); continue
    if s.startswith("- "):
        inline(doc.add_paragraph(style="List Bullet"), s[2:]); continue
    m = re.match(r"^(\d+)\.\s+(.*)", s)
    if m:
        inline(doc.add_paragraph(style="List Number"), m.group(2)); continue
    inline(doc.add_paragraph(), s)

doc.save(OUT)
print("wrote", OUT, "| paragraphs:", len(doc.paragraphs))
