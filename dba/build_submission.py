"""
Build the submission packet:
  1. Proceeding_Email_and_Notes_2026-06-01.docx  (from the .md)
  2. Data_Collection_Readiness.docx              (from the .md, professor-friendly)
  3. YMalik_DBA_Submission_2026-06-01.zip        (curated upload set)
"""
import re, zipfile, os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def new_doc():
    doc = Document()
    n = doc.styles["Normal"]
    n.font.name = "Times New Roman"; n.font.size = Pt(11)
    n._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    n.paragraph_format.line_spacing = 1.15
    n.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    n.paragraph_format.space_after = Pt(4)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.9)
        s.left_margin = s.right_margin = Inches(1)
    return doc

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), fill)
    tcPr.append(sh)

def inline(p, text):
    for tok in re.split(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)", text):
        if not tok: continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            r = p.add_run(tok[1:-1]); r.italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(10)
        else:
            p.add_run(tok)

def render_md(doc, md, drop_front=True):
    lines = md.split("\n"); i = 0
    if drop_front and lines and lines[0].strip() == "---":
        i = 1
        while i < len(lines) and lines[i].strip() != "---": i += 1
        i += 1
    n = len(lines)
    while i < n:
        s = lines[i].rstrip().strip()
        if s == "" or s.startswith("<!--"): i += 1; continue
        if s == "---": doc.add_page_break(); i += 1; continue
        if s.startswith("#### "): doc.add_heading(s[5:], 4); i += 1; continue
        if s.startswith("### "):  doc.add_heading(s[4:], 3); i += 1; continue
        if s.startswith("## "):   doc.add_heading(s[3:], 2); i += 1; continue
        if s.startswith("# "):    doc.add_heading(s[2:], 1); i += 1; continue
        if s.startswith("- ") or s.startswith("* "):
            p = doc.add_paragraph(style="List Bullet"); inline(p, s[2:]); i += 1; continue
        mnum = re.match(r"^(\d+)\.\s+(.*)", s)
        if mnum:
            p = doc.add_paragraph(style="List Number"); inline(p, mnum.group(2)); i += 1; continue
        if "|" in s and i+1 < n and re.match(r"^\s*\|?\s*[-:|\s]+\|", lines[i+1]):
            rows = []
            while i < n and "|" in lines[i].strip():
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells); i += 1
                if i < n and re.match(r"^\s*\|?\s*[-:|\s]+\|", lines[i]): i += 1
            if rows:
                cols = max(len(r) for r in rows)
                t = doc.add_table(rows=len(rows), cols=cols); t.style = "Light Grid Accent 1"
                for ri, row in enumerate(rows):
                    for ci in range(cols):
                        c = t.cell(ri, ci); c.text = ""
                        pp = c.paragraphs[0]; pp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        inline(pp, row[ci] if ci < len(row) else "")
                doc.add_paragraph().paragraph_format.space_after = Pt(0)
            continue
        p = doc.add_paragraph(); inline(p, s); i += 1

# 1. Proceeding email + notes
d1 = new_doc()
render_md(d1, open("Proceeding_Email_and_Notes_2026-06-01.md", encoding="utf-8").read())
d1.save("Proceeding_Email_and_Notes_2026-06-01.docx")
print("wrote Proceeding_Email_and_Notes_2026-06-01.docx")

# 2. Readiness checklist as docx
d2 = new_doc()
d2.add_heading("Data Collection Readiness — Summary for Advisor", 1)
render_md(d2, open("Data_Collection_Readiness.md", encoding="utf-8").read())
d2.save("Data_Collection_Readiness.docx")
print("wrote Data_Collection_Readiness.docx")

# 3. ZIP the curated upload set
zip_items = [
    "Research_Paper_YMalik_v4.docx",
    "ALL_IN_ONE_DBA_Package_v3.docx",
    "Proceeding_Email_and_Notes_2026-06-01.docx",
    "Data_Collection_Readiness.docx",
    "LinkedIn_Outreach_Pack.docx",
    "03_Recruitment_and_Pilot/YMalik_Informed_Pilot_Protocol_READY_2026-06-01.docx",
    "03_Recruitment_and_Pilot/YMalik_Pilot_Feedback_Form_and_Revision_Log_READY_2026-06-01.xlsx",
    "03_Recruitment_and_Pilot/YMalik_CloudResearch_Launch_Draft_READY_2026-06-01.docx",
    "03_Recruitment_and_Pilot/YMalik_MTurk_Backup_Launch_Draft_READY_2026-06-01.docx",
]
zname = "YMalik_DBA_Submission_2026-06-01.zip"
with zipfile.ZipFile(zname, "w", zipfile.ZIP_DEFLATED) as z:
    for item in zip_items:
        if os.path.exists(item):
            # flatten the recruitment-folder files into the zip root for easy attaching
            arc = os.path.basename(item)
            z.write(item, arc)
            print("  + zipped", arc)
        else:
            print("  ! MISSING", item)
print("wrote", zname)
print("zip size:", os.path.getsize(zname), "bytes")
