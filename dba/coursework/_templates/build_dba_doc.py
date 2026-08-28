#!/usr/bin/env python3
"""
DBA coursework document builder — Yasir A. Malik.

House format taken from his own submitted work, not invented here:
  dba/Research_Paper_YMalik_SUBMISSION.docx
  dba/Data_Collection_Readiness.docx
  dba/2026-06-14_Catchup_Status_to_Rey.docx

  - Times New Roman 11pt, 1.15 line spacing
  - Native Word Heading 1 / Heading 2 / List Bullet styles, so the document
    opens with a working navigation pane and the professor's own styling
  - An identification block under the title: project, author, programme,
    course, instructor, date
  - NO wordmark, NO repository URL, NO "Audit the Algorithm"

That last line is the point of this file existing separately from
build_submission.py. The consulting brand belongs on portfolio and public
documents. A class submission carries the university's identity, not a
practice's — a professor reading a consulting letterhead on coursework sees a
vendor, not a doctoral student.

Usage:  python3 build_dba_doc.py <spec.json> <out.docx>
"""
import json, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = 'Times New Roman'
INK  = RGBColor(0x00, 0x00, 0x00)

def _font(obj, name=FONT, size=11, bold=None, italic=None, color=INK,
          highlight=False):
    obj.font.name = name
    obj.font.size = Pt(size)
    if color is not None: obj.font.color.rgb = color
    if bold is not None: obj.bold = bold
    if italic is not None: obj.italic = italic
    if highlight: obj.font.highlight_color = WD_COLOR_INDEX.YELLOW
    rpr = obj._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rf.set(qn(a), name)
    return obj

def _segs(par, segs, size=11):
    """A paragraph body: a string, or a list of strings and {t,b,i,hl} dicts."""
    if isinstance(segs, str): segs = [segs]
    for s in segs:
        if isinstance(s, str):
            _font(par.add_run(s), size=size)
        else:
            _font(par.add_run(s["t"]), size=size,
                  bold=s.get("b", False), italic=s.get("i", False),
                  highlight=s.get("hl", False))

def build(spec, out):
    doc = Document()

    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(1.0)
    sec.left_margin = sec.right_margin = Inches(1.0)

    ds = spec.get("double_spaced", False)
    size, ls = (12, 2.0) if ds else (11, 1.15)

    n = doc.styles['Normal']
    _font(n, size=size)
    n.paragraph_format.line_spacing = ls
    n.paragraph_format.space_after = Pt(8)

    for sname, ssize in (('Heading 1', size + 3), ('Heading 2', size + 1)):
        st = doc.styles[sname]
        _font(st, size=ssize, bold=True, color=INK)
        st.paragraph_format.space_before = Pt(14)
        st.paragraph_format.space_after = Pt(6)
    _font(doc.styles['List Bullet'], size=size)

    # ---- title ---------------------------------------------------------
    h = doc.add_paragraph(spec["title"], style='Heading 1')
    for r in h.runs: _font(r, size=size + 4, bold=True)

    if spec.get("subtitle"):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10)
        _font(p.add_run(spec["subtitle"]), size=size + 1, italic=True)

    # ---- identification block ------------------------------------------
    for line in spec.get("ident", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        _font(p.add_run(line), size=size - 0.5)
    if spec.get("ident"):
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ---- body ------------------------------------------------------------
    def emit(items):
        for it in items:
            if isinstance(it, dict) and "h1" in it:
                p = doc.add_paragraph(it["h1"], style='Heading 1')
                for r in p.runs: _font(r, size=size + 3, bold=True)
            elif isinstance(it, dict) and "h2" in it:
                p = doc.add_paragraph(it["h2"], style='Heading 2')
                for r in p.runs: _font(r, size=size + 1, bold=True)
            elif isinstance(it, dict) and "bullet" in it:
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.line_spacing = ls
                _segs(p, it["bullet"], size=size)
            elif isinstance(it, dict) and "pagebreak" in it:
                doc.add_page_break()
            elif isinstance(it, dict) and "__table__" in it:
                rows = spec["tables"][it["__table__"]]
                t = doc.add_table(rows=0, cols=len(rows[0]))
                t.style = 'Table Grid'
                for ri, row in enumerate(rows):
                    cells = t.add_row().cells
                    for ci, val in enumerate(row):
                        cp = cells[ci].paragraphs[0]
                        cp.paragraph_format.space_after = Pt(2)
                        cp.paragraph_format.line_spacing = 1.0
                        _font(cp.add_run(val), size=size - 1.5,
                              bold=(ri == 0), color=INK)
                doc.add_paragraph().paragraph_format.space_after = Pt(6)
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing = ls
                p.paragraph_format.space_after = Pt(8)
                _segs(p, it, size=size)

    emit(spec["body"])

    for tb in (spec.get("table"), spec.get("table2")):
        if not tb: continue
        if tb.get("caption"):
            p = doc.add_paragraph(tb["caption"], style='Heading 2')
            for r in p.runs: _font(r, size=size + 1, bold=True)
        t = doc.add_table(rows=0, cols=len(tb["rows"][0]))
        t.style = 'Table Grid'
        for ri, row in enumerate(tb["rows"]):
            cells = t.add_row().cells
            for ci, val in enumerate(row):
                cp = cells[ci].paragraphs[0]
                cp.paragraph_format.space_after = Pt(2)
                cp.paragraph_format.line_spacing = 1.0
                _font(cp.add_run(val), size=size - 1.5,
                      bold=(ri == 0), color=INK)

    if spec.get("appendix"):
        doc.add_page_break()
        emit(spec["appendix"])

    doc.save(out)
    print("wrote", out)

if __name__ == "__main__":
    build(json.load(open(sys.argv[1])), sys.argv[2])
