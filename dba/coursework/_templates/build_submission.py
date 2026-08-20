#!/usr/bin/env python3
"""
House Word-document builder for FIU DBA submissions — Yasir A. Malik.

Design comes from live.html: accent #1F4E79, ink #14171C, muted #767D86,
Georgia/Palatino serif body. Wordmark is drawn as text (the logo.svg is itself
pure text) so the .docx stays self-contained with no image dependency.

Anything the author must supply is written in ALL CAPS with a YELLOW HIGHLIGHT
so it cannot be missed or accidentally submitted.

Usage:  python3 build_submission.py <spec.json> <out.docx>
"""
import json, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
INK    = RGBColor(0x14, 0x17, 0x1C)
MUTED  = RGBColor(0x76, 0x7D, 0x86)
GOLD   = RGBColor(0xB8, 0x86, 0x0B)

REPO = "github.com/AuditingAI/Profile"
SITE = "auditingai.github.io"

def rule(par, color="1F4E79", size=12):
    p = par._p.get_or_add_pPr()
    bdr = OxmlElement('w:pBdr'); bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), str(size))
    bot.set(qn('w:space'), '4'); bot.set(qn('w:color'), color)
    bdr.append(bot); p.append(bdr)

def run(par, text, *, font="Georgia", size=11.5, color=INK, bold=False,
        italic=False, caps=False, spacing=None, highlight=False):
    r = par.add_run(text)
    r.font.name = font; r.font.size = Pt(size); r.font.color.rgb = color
    r.bold = bold; r.italic = italic
    if caps: r.font.all_caps = True
    if highlight: r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    rpr = r._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    for a in ('w:ascii','w:hAnsi','w:cs'): rf.set(qn(a), font)
    if spacing:
        sp = OxmlElement('w:spacing'); sp.set(qn('w:val'), str(int(spacing*20)))
        rpr.append(sp)
    return r

def build(spec, out):
    doc = Document()
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = Inches(0.85)
    s.left_margin = s.right_margin = Inches(1.0)

    n = doc.styles['Normal']
    n.font.name = 'Georgia'; n.font.size = Pt(11.5); n.font.color.rgb = INK
    n.paragraph_format.space_after = Pt(9)
    n.paragraph_format.line_spacing = 1.15

    # ---- wordmark ------------------------------------------------------
    wm = doc.add_paragraph(); wm.paragraph_format.space_after = Pt(2)
    run(wm, "Audit ", font="Georgia", size=15, color=ACCENT, bold=True)
    run(wm, "the ",   font="Georgia", size=12, color=MUTED, italic=True)
    run(wm, "Algorithm", font="Georgia", size=15, color=ACCENT, bold=True)

    # ---- repo / site line ----------------------------------------------
    lk = doc.add_paragraph(); lk.paragraph_format.space_after = Pt(9)
    run(lk, f"{REPO}   ·   {SITE}", font="Consolas", size=7.5,
        color=MUTED, spacing=0.9)
    rule(lk)

    # ---- title + meta ---------------------------------------------------
    t = doc.add_paragraph(); t.paragraph_format.space_before = Pt(10)
    t.paragraph_format.space_after = Pt(2)
    run(t, spec["title"], font="Georgia", size=15, color=INK, bold=True)

    m = doc.add_paragraph(); m.paragraph_format.space_after = Pt(16)
    run(m, spec["meta"], font="Consolas", size=8, color=MUTED,
        caps=True, spacing=1.1)

    # ---- body ------------------------------------------------------------
    for para in spec["body"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(9)
        p.paragraph_format.line_spacing = 1.15
        for seg in para:
            if isinstance(seg, str):
                run(p, seg)
            else:
                run(p, seg["t"], italic=seg.get("i", False),
                    bold=seg.get("b", False),
                    caps=seg.get("caps", False),
                    highlight=seg.get("hl", False),
                    color=ACCENT if seg.get("hl") else INK)

    # ---- footer blocks ----------------------------------------------------
    for blk in spec.get("blocks", []):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(3)
        run(h, blk["heading"], font="Consolas", size=8, color=MUTED,
            caps=True, spacing=1.1, bold=True)
        rule(h, color="D8D4CB", size=6)
        for para in blk["body"]:
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.line_spacing = 1.1
            for seg in para:
                if isinstance(seg, str):
                    run(p, seg, size=9, color=MUTED)
                else:
                    run(p, seg["t"], size=9,
                        color=ACCENT if seg.get("hl") else MUTED,
                        bold=seg.get("b", False), italic=seg.get("i", False),
                        caps=seg.get("caps", False), highlight=seg.get("hl", False))

    doc.save(out)
    print("wrote", out)

if __name__ == "__main__":
    build(json.load(open(sys.argv[1])), sys.argv[2])
